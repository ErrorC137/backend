"""In-memory document ingestion — no disk persistence."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

from PyPDF2 import PdfReader
from docx import Document


class DocumentExtractionError(ValueError):
    """Raised when an upload does not contain usable human-readable text."""


SECTION_PATTERNS = {
    "abstract": r"(?is)(?:^|\n)\s*(?:abstract|summary)\s*[:\-]?\s*(.+?)(?=\n\s*(?:introduction|1[\.\)]|keywords|background|index terms)\b)",
    "introduction": r"(?is)(?:^|\n)\s*(?:1[\.\)]?\s*)?introduction\s*[:\-]?\s*(.+?)(?=\n\s*(?:2[\.\)]|related work|methods?|materials)\b)",
    "methods": r"(?is)(?:^|\n)\s*(?:2[\.\)]?\s*)?(?:methods?|methodology|materials and methods|experimental)\s*[:\-]?\s*(.+?)(?=\n\s*(?:3[\.\)]|results?|findings|discussion)\b)",
    "results": r"(?is)(?:^|\n)\s*(?:3[\.\)]?\s*)?(?:results?|findings)\s*[:\-]?\s*(.+?)(?=\n\s*(?:4[\.\)]|discussion|conclusion|references)\b)",
    "conclusion": r"(?is)(?:^|\n)\s*(?:4[\.\)]?\s*)?(?:conclusions?|discussion)\s*[:\-]?\s*(.+?)(?=\n\s*(?:references|acknowledg|appendix|bibliography)\b|$)",
}


@dataclass
class ParsedDocument:
    raw_text: str
    abstract: str
    methodology: str
    claims_outcomes: str
    document_type: str = "scientific_paper"
    parsing_confidence: float = 0.5
    sections_found: list[str] = field(default_factory=list)
    authors: list[str] = field(default_factory=list)
    institutions: list[str] = field(default_factory=list)
    affiliations: list[str] = field(default_factory=list)
    keywords: list[str] = field(default_factory=list)
    references: list[str] = field(default_factory=list)
    team_expertise_score: float = 0.5
    institution_reputation_score: float = 0.5


def _strip_metadata(text: str) -> str:
    text = re.sub(r"(?i)(corresponding author|affiliation|university of)[^\n]{0,120}", "", text)
    text = re.sub(r"(?i)(email|@)[^\s]{3,80}", "", text)
    text = re.sub(r"(?i)^\s*arxiv:[^\n]+\n", "", text)
    return text


def _looks_like_pdf_structure(text: str) -> bool:
    """Identify PDF/XML internals before they are scored as research content."""
    markers = re.findall(
        r"(?i)(?:\b(?:endobj|startxref|xref|pdfaProperty|rdf:|dc:identifier)\b|Parent\s+\d+\s+\d+\s+R|/Type\s*/|/Title\()",
        text,
    )
    words = re.findall(r"[A-Za-z]{3,}", text)
    return len(markers) >= 3 or (len(text) > 300 and len(words) < len(text) / 30)


def _extract_section(text: str, pattern: str) -> str | None:
    match = re.search(pattern, text)
    return match.group(1).strip() if match else None


def _extract_authors(text: str) -> list[str]:
    """Extract author names from document text."""
    authors = []
    # Common author patterns
    patterns = [
        r"(?i)(?:authors?|by)[\s:]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,2}(?:,\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,2})*)",
        r"(?i)^([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2}),\s*[A-Z][a-z]+",
        r"(?i)([A-Z][a-z]+,\s*[A-Z][a-z]+\s+and\s+[A-Z][a-z]+)",
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text[:5000])
        authors.extend(matches[:10])
    # Clean and deduplicate
    authors = list(set([a.strip() for a in authors if len(a.strip()) > 2]))
    return authors[:10]


def _extract_institutions(text: str) -> list[str]:
    """Extract institution names from document text."""
    institutions = []
    # Common institution patterns
    patterns = [
        r"(?i)(?:university|institute|laboratory|college|school)[\s,]+(?:of|for)?\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)",
        r"(?i)([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,2}\s+(?:University|Institute|Laboratory|College))",
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text[:5000])
        institutions.extend(matches)
    # Clean and deduplicate
    institutions = list(set([i.strip() for i in institutions if len(i.strip()) > 3]))
    return institutions[:5]


def _assess_institution_reputation(institutions: list[str]) -> float:
    """Assess institution reputation based on keywords."""
    if not institutions:
        return 0.3  # Low score if no institutions found
    
    reputation_keywords = {
        "high": ["mit", "stanford", "harvard", "caltech", "oxford", "cambridge", "eth", "imperial", "national", "federal", "institute"],
        "medium": ["university", "college", "laboratory", "research", "technology", "science"],
    }
    
    institution_text = " ".join(institutions).lower()
    score = 0.5  # Base score
    
    for keyword in reputation_keywords["high"]:
        if keyword in institution_text:
            score += 0.15
    
    for keyword in reputation_keywords["medium"]:
        if keyword in institution_text:
            score += 0.05
    
    return min(1.0, score)


def _assess_author_expertise(authors: list[str], institutions: list[str]) -> float:
    """Assess author expertise based on count and affiliations."""
    if not authors:
        return 0.2  # Low score if no authors found
    
    # More authors generally indicates larger team/collaboration
    author_score = min(0.6, len(authors) * 0.1)
    
    # Institution reputation contributes to expertise assessment
    institution_score = _assess_institution_reputation(institutions) * 0.4
    
    return min(1.0, author_score + institution_score)


def _extract_keywords(text: str) -> list[str]:
    """Extract keywords from document text."""
    keywords = []
    # Look for keywords section
    keyword_pattern = r"(?i)(?:keywords?|index terms)[\s:]+(.+?)(?:\n|$|abstract|introduction)"
    match = re.search(keyword_pattern, text[:3000])
    if match:
        keyword_text = match.group(1)
        # Split by common delimiters
        keywords = re.split(r"[,;·•\n]", keyword_text)
    else:
        # Extract technical terms as fallback
        technical_terms = [
            "synthesis", "fabrication", "characterization", "optimization", "analysis",
            "performance", "efficiency", "stability", "durability", "scalability",
            "materials", "composite", "nanoparticle", "polymer", "ceramic", "alloy",
            "processing", "manufacturing", "testing", "validation", "evaluation",
            "properties", "structure", "composition", "morphology", "surface",
            "energy", "storage", "conversion", "generation", "transmission",
            "carbon", "capture", "sequestration", "reduction", "utilization",
            "catalyst", "reaction", "mechanism", "kinetics", "thermodynamics",
            "electronic", "optical", "magnetic", "mechanical", "thermal",
            "device", "system", "application", "implementation", "integration"
        ]
        text_lower = text.lower()
        keywords = [term for term in technical_terms if term in text_lower]
    
    # Clean and deduplicate
    keywords = list(set([k.strip() for k in keywords if len(k.strip()) > 2]))
    return keywords[:15]


def _extract_references(text: str) -> list[str]:
    """Extract reference citations from document text."""
    references = []
    # Look for references section
    ref_pattern = r"(?i)(?:references|bibliography)[\s:]+(.+)"
    match = re.search(ref_pattern, text[-10000:])
    if match:
        ref_text = match.group(1)
        # Extract individual references
        ref_matches = re.findall(r"\[\d+\][^\[\]]{20,200}", ref_text)
        references = ref_matches[:20]
    return references


def _infer_document_type(text: str, filename: str) -> str:
    """Enhanced document type inference for IP-specific structures."""
    lower = text[:5000].lower()
    name = filename.lower()
    
    # Patent-specific patterns
    if "patent" in name or re.search(r"(?i)claim\s+\d+\.", text[:5000]):
        return "patent_draft"
    if re.search(r"(?i)(independent claim|dependent claim|what is claimed)", text[:5000]):
        return "patent_application"
    
    # Preprint patterns
    if "preprint" in lower or "arxiv" in lower or "biorxiv" in lower:
        return "preprint"
    
    # Technical report patterns
    if re.search(r"(?i)(technical report|white paper|working paper)", text[:2000]):
        return "technical_report"
    
    # Invention disclosure patterns
    if re.search(r"(?i)(invention disclosure|disclosure form|inventor disclosure)", text[:3000]):
        return "invention_disclosure"
    
    # Research grant proposal patterns
    if re.search(r"(?i)(grant proposal|research proposal|funding application)", text[:3000]):
        return "grant_proposal"
    
    # Default to scientific paper
    return "scientific_paper"


def _fallback_chunk_extraction(text: str) -> tuple[str, str, str]:
    """Handle papers without standard section headers (common in preprints/exports)."""
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if len(p.strip()) > 40]
    if not paragraphs:
        chunk = text.strip()
        third = max(len(chunk) // 3, 1)
        return chunk[:third], chunk[third : third * 2], chunk[third * 2 :]

    n = len(paragraphs)
    abstract = paragraphs[0]
    if n >= 4:
        methodology = "\n\n".join(paragraphs[1 : n // 2])
        claims = "\n\n".join(paragraphs[n // 2 :])
    elif n == 3:
        methodology = paragraphs[1]
        claims = paragraphs[2]
    else:
        methodology = paragraphs[1] if n > 1 else paragraphs[0]
        claims = paragraphs[-1]
    return abstract, methodology, claims


def _extract_sections(text: str) -> tuple[str, str, str, list[str], float]:
    found: list[str] = []
    abstract = _extract_section(text, SECTION_PATTERNS["abstract"])
    if abstract:
        found.append("abstract")

    methods = _extract_section(text, SECTION_PATTERNS["methods"])
    if methods:
        found.append("methods")

    results = _extract_section(text, SECTION_PATTERNS["results"])
    conclusion = _extract_section(text, SECTION_PATTERNS["conclusion"])
    claims_parts = [p for p in (results, conclusion) if p]
    claims = "\n\n".join(claims_parts) if claims_parts else None

    if results:
        found.append("results")
    if conclusion:
        found.append("conclusion")

    if abstract and methods and claims:
        confidence = 0.95
    elif abstract and (methods or claims):
        confidence = 0.75
    elif abstract:
        confidence = 0.55
        abstract, methods, claims = _fallback_chunk_extraction(text)
        found.append("fallback_chunking")
    else:
        abstract, methods, claims = _fallback_chunk_extraction(text)
        found = ["fallback_chunking"]
        confidence = 0.45

    return abstract, methods or text[800:3500].strip(), claims or text[-2000:].strip(), found, confidence


def parse_upload(filename: str, content: bytes) -> ParsedDocument:
    name = filename.lower()
    if name.endswith(".pdf"):
        raw = ""
        extraction_method = ""
        
        try:
            # Method 1: Try PyMuPDF (fitz) - most robust
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(stream=content, filetype="pdf")
                pages_text = []
                for page in doc:
                    try:
                        text = page.get_text()
                        if text and len(text.strip()) > 10:
                            pages_text.append(text)
                    except Exception:
                        continue
                
                if pages_text:
                    raw = "\n".join(pages_text)
                    extraction_method = "PyMuPDF"
                    doc.close()
                else:
                    raise ImportError("PyMuPDF extraction returned no text")
            except ImportError:
                extraction_method = "PyMuPDF not available"
                pass
            
            # Method 2: Try pdfplumber
            if not raw or len(raw.strip()) < 50:
                try:
                    import pdfplumber
                    with pdfplumber.open(io.BytesIO(content)) as pdf:
                        pages = []
                        for page in pdf.pages:
                            try:
                                text = page.extract_text()
                                if text and len(text.strip()) > 10:
                                    cleaned = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
                                    if len(cleaned.strip()) > 10:
                                        pages.append(cleaned)
                            except Exception:
                                continue
                        
                        if pages:
                            raw = "\n".join(pages)
                            extraction_method = "pdfplumber"
                        else:
                            raise ImportError("pdfplumber extraction failed")
                except ImportError:
                    extraction_method = "pdfplumber not available"
                    pass
            
            # Method 3: Try PyPDF2 with enhanced cleaning
            if not raw or len(raw.strip()) < 50:
                try:
                    reader = PdfReader(io.BytesIO(content))
                    pages = []
                    for page in reader.pages:
                        try:
                            text = page.extract_text()
                            if text and len(text.strip()) > 10:
                                # Aggressive cleaning
                                cleaned = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
                                # Remove PDF structure artifacts
                                cleaned = re.sub(r'\s*\d+\s*\.\s*\d+\s*', '', cleaned)
                                cleaned = re.sub(r'\s*endobj\s*', '', cleaned)
                                cleaned = re.sub(r'\s*stream\s*', '', cleaned)
                                cleaned = re.sub(r'\s*endstream\s*', '', cleaned)
                                cleaned = re.sub(r'\s*xref\s*', '', cleaned)
                                cleaned = re.sub(r'\s*startxref\s*', '', cleaned)
                                cleaned = re.sub(r'\s*<<\s*/\s*\w+\s*/\s*\w+\s*>>\s*', '', cleaned)
                                cleaned = re.sub(r'\s*/Type\s*/\w+\s*', '', cleaned)
                                cleaned = re.sub(r'\s/Subtype\s*/\w+\s*', '', cleaned)
                                cleaned = re.sub(r'\s/Rect\s*\[.*?\]\s*', '', cleaned)
                                if len(cleaned.strip()) > 10:
                                    pages.append(cleaned)
                        except Exception:
                            continue
                    
                    if pages:
                        raw = "\n".join(pages)
                        extraction_method = "PyPDF2"
                except Exception as e:
                    extraction_method = f"PyPDF2 failed: {str(e)}"
            
            # Method 4: OCR fallback for image-based PDFs
            if not raw or len(raw.strip()) < 50:
                try:
                    import pytesseract
                    from PIL import Image
                    import io
                    
                    # Try PyMuPDF to convert pages to images for OCR
                    try:
                        import fitz
                        doc = fitz.open(stream=content, filetype="pdf")
                        pages_text = []
                        for page in doc:
                            try:
                                # Convert page to image
                                pix = page.get_pixmap()
                                img_data = pix.tobytes("png")
                                image = Image.open(io.BytesIO(img_data))
                                
                                # OCR the image
                                text = pytesseract.image_to_string(image)
                                if text and len(text.strip()) > 10:
                                    pages_text.append(text)
                            except Exception:
                                continue
                        
                        if pages_text:
                            raw = "\n".join(pages_text)
                            extraction_method = "OCR"
                        doc.close()
                    except Exception as e:
                        extraction_method = f"OCR failed: {str(e)}"
                except ImportError:
                    extraction_method = "OCR not available (requires pytesseract and Pillow)"
                    pass
            
            # Never decode raw PDF bytes as a last resort. It produces PDF/XMP
            # object syntax which is not document text and must not be analyzed.
            if not raw or len(raw.strip()) < 50:
                raise DocumentExtractionError(
                    "No usable scientific text could be extracted from this PDF. "
                    "Upload a text-based PDF/DOCX, or a PDF with an OCR text layer."
                )
                try:  # pragma: no cover - retained only as historical reference
                    raw = content.decode('utf-8', errors='ignore')
                    # Comprehensive PDF artifact removal - more aggressive
                    raw = re.sub(r'[^\x20-\x7E\n\r\t]', '', raw)
                    # Remove all PDF structure keywords
                    pdf_keywords = ['endobj', 'stream', 'endstream', 'xref', 'startxref', 'obj', 'R']
                    for keyword in pdf_keywords:
                        raw = re.sub(rf'\b{keyword}\b', '', raw, flags=re.IGNORECASE)
                    # Remove PDF dictionary structures
                    raw = re.sub(r'<<.*?>>', '', raw, flags=re.DOTALL)
                    # Remove PDF array structures
                    raw = re.sub(r'\[.*?\]', '', raw, flags=re.DOTALL)
                    # Remove PDF name objects
                    raw = re.sub(r'/\w+', '', raw)
                    # Remove PDF numeric references
                    raw = re.sub(r'\b\d+\s+\d+\s+R\b', '', raw)
                    raw = re.sub(r'\b\d+\s+obj\b', '', raw)
                    # Remove common PDF artifacts
                    raw = re.sub(r'/Type|/Subtype|/Rect|/Action|/Dest|/Parent|/First|/Last|/Count|/Title|/Prev|/Next|/StructParent|/F|/BS|/S|/W|/CA|/ca|/LW|/Filter|/Length|/BitsPerComponent|/ColorSpace|/Width|/Height|/ICCBased|/Separation', '', raw, flags=re.IGNORECASE)
                    # Remove hex-encoded content
                    raw = re.sub(r'[0-9A-Fa-f]{20,}', '', raw)
                    # Remove base64-like content
                    raw = re.sub(r'[A-Za-z0-9+/]{50,}={0,2}', '', raw)
                    # Remove repeated special characters
                    raw = re.sub(r'[<>{}\\]+', '', raw)
                    # Clean up whitespace
                    raw = re.sub(r'\s+', ' ', raw)
                    extraction_method = "raw_decode"
                except Exception as e:
                    raw = f"PDF parsing error: {str(e)}. Method used: {extraction_method}"
            
            # Final validation
            if not raw or len(raw.strip()) < 50:
                raw = f"Document content could not be properly extracted. The PDF may be image-based or corrupted. Extraction method: {extraction_method}. Please try a different file format."
            
        except DocumentExtractionError:
            raise
        except Exception as e:
            raise DocumentExtractionError(f"PDF extraction failed: {e}") from e
            
    elif name.endswith(".docx"):
        doc = Document(io.BytesIO(content))
        raw = "\n".join(p.text for p in doc.paragraphs)
    else:
        raw = content.decode("utf-8", errors="replace")

    # Clean the extracted text - aggressive PDF artifact removal
    raw = _strip_metadata(raw)
    
    # Remove all non-ASCII characters first
    raw = re.sub(r'[^\x20-\x7E\n\r\t]', '', raw)
    
    # Aggressive PDF artifact removal
    raw = re.sub(r'/\w+', '', raw)  # Remove all PDF name objects
    raw = re.sub(r'\b\d+\s+\d+\s+R\b', '', raw)  # Remove PDF references
    raw = re.sub(r'\b\d+\s+obj\b', '', raw)  # Remove object references
    raw = re.sub(r'<<.*?>>', '', raw, flags=re.DOTALL)  # Remove PDF dictionaries
    raw = re.sub(r'\[.*?\]', '', raw, flags=re.DOTALL)  # Remove PDF arrays
    raw = re.sub(r'\bendobj\b|\bstream\b|\bendstream\b|\bxref\b|\bstartxref\b', '', raw, flags=re.IGNORECASE)
    raw = re.sub(r'/Type|/Subtype|/Rect|/Action|/Dest|/Parent|/First|/Last|/Count|/Title|/Prev|/Next|/StructParent|/F|/BS|/S|/W|/CA|/ca|/LW|/Filter|/Length|/BitsPerComponent|/ColorSpace|/Width|/Height|/ICCBased|/Separation', '', raw, flags=re.IGNORECASE)
    raw = re.sub(r'/Pg\s+\d+\s+\d+\s+R', '', raw)  # Remove page references
    raw = re.sub(r'/K\s*\[.*?\]', '', raw, flags=re.DOTALL)  # Remove PDF arrays
    raw = re.sub(r'/H\d+', '', raw)  # Remove heading references
    raw = re.sub(r'/P\s+\d+\s+\d+\s+R', '', raw)  # Remove paragraph references
    raw = re.sub(r'/T\s*\([^)]*\)', '', raw)  # Remove text references
    raw = re.sub(r'/E\s*\([^)]*\)', '', raw)  # Remove element references
    raw = re.sub(r'/A\s*\[.*?\]', '', raw, flags=re.DOTALL)  # Remove action arrays
    raw = re.sub(r'[0-9A-Fa-f]{20,}', '', raw)  # Remove hex-encoded content
    raw = re.sub(r'[A-Za-z0-9+/]{50,}={0,2}', '', raw)  # Remove base64-like content
    raw = re.sub(r'[<>{}\\]+', '', raw)  # Remove special characters
    raw = re.sub(r'\s+', ' ', raw)  # Clean up whitespace
    
    # Never score parser errors or PDF syntax as research content.
    if len(raw.strip()) < 50 or _looks_like_pdf_structure(raw):
        raise DocumentExtractionError(
            "No usable scientific text could be extracted. Upload a text-based PDF, DOCX, or plain-text file."
        )
    
    doc_type = _infer_document_type(raw, filename)
    abstract, methodology, claims, sections_found, confidence = _extract_sections(raw)
    
    # Extract additional metadata
    authors = _extract_authors(raw)
    institutions = _extract_institutions(raw)
    keywords = _extract_keywords(raw)
    references = _extract_references(raw)
    
    # Assess team expertise based on authors and institutions
    team_expertise_score = _assess_author_expertise(authors, institutions)
    institution_reputation_score = _assess_institution_reputation(institutions)
    
    return ParsedDocument(
        raw_text=raw,
        abstract=abstract,
        methodology=methodology,
        claims_outcomes=claims,
        document_type=doc_type,
        parsing_confidence=confidence,
        sections_found=sections_found,
        authors=authors,
        institutions=institutions,
        affiliations=institutions,  # Same as institutions for now
        keywords=keywords,
        references=references,
        team_expertise_score=team_expertise_score,
        institution_reputation_score=institution_reputation_score,
    )
